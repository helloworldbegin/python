from docx import Document
document=Document()
document.add_heading('中华人民共和国', 0)
p = document.add_paragraph('中华人民共和国成立于1949年！')
document.add_page_break()
document.save('c:/tmp/demo.docx')